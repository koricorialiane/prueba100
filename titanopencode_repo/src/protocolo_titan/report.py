import re
from datetime import date
from html import escape as html_escape
from pathlib import Path

import markdown
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _format_number(value: float, digits: int = 4) -> str:
    if digits == 0:
        return str(int(round(value)))
    return f"{value:.{digits}f}".rstrip("0").rstrip(".")


def _format_markdown_cell(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return format(value, "g")
    return str(value).replace("|", "\\|")


def _dataframe_to_markdown(df: pd.DataFrame) -> str:
    columns = [str(column) for column in df.columns]
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]

    for row in df.itertuples(index=False, name=None):
        lines.append("| " + " | ".join(_format_markdown_cell(value) for value in row) + " |")

    return "\n".join(lines)


def _append_section(lines: list[str], title: str, body: list[str]) -> None:
    lines.append(title)
    lines.append("")
    lines.extend(body)
    lines.append("")


def _mobility_case_lines(row: pd.Series) -> list[str]:
    speed_kmh = float(row["speed_kmh"])
    speed_ms = float(row["speed_ms"])
    fd = float(row["max_doppler_hz"])
    tc_ms = float(row["coherence_time_ms"])
    tc_s = tc_ms / 1e3
    timeslot_ms = float(row["gsm_timeslot_ms"])
    ratio = float(row["coherence_to_timeslot_ratio"])
    stability = str(row["stability_class"])
    fc_mhz = float(row["carrier_frequency_mhz"])

    return [
        f"#### Caso {speed_kmh:g} km/h",
        "",
        f"1. Conversión de velocidad: `v = {speed_kmh:g} / 3.6 = {_format_number(speed_ms)} m/s`.",
        (
            "2. Desviación Doppler máxima: "
            f"`f_d = v · f_c / c = {_format_number(speed_ms)} · {fc_mhz:g}e6 / 3e8 = {_format_number(fd)} Hz`."
        ),
        (
            "3. Tiempo de coherencia: "
            f"`T_c ≈ 0.423 / f_d = 0.423 / {_format_number(fd)} = {_format_number(tc_s, 6)} s = {_format_number(tc_ms)} ms`."
        ),
        (
            "4. Comparación con el timeslot GSM: "
            f"`T_c / T_slot = {_format_number(tc_ms)} / {_format_number(timeslot_ms, 3)} = {_format_number(ratio)}`."
        ),
        (
            f"Interpretación: el canal se clasifica como **{stability}**. "
            "Por tanto, la ráfaga puede considerarse cuasiestática o estable dentro del timeslot, "
            "pero la robustez del enlace sigue dependiendo de codificación, entrelazado y margen de diseño."
        ),
        "",
        "Mini conclusión parcial: la movilidad extrema incrementa el Doppler y reduce el tiempo de coherencia; sin embargo, con los parámetros base el canal sigue siendo manejable dentro de una ráfaga GSM.",
    ]


def _noise_case_lines(row: pd.Series) -> str:
    rbw_hz = float(row["rbw_hz"])
    rbw_khz = float(row["rbw_khz"])
    nf_db = float(row["noise_figure_db"])
    noise_floor_dbm = float(row["noise_floor_dbm"])
    return (
        f"- RBW = {rbw_khz:g} kHz: `N = -174 + 10 log10({rbw_hz:.0f}) + {nf_db:g} = {_format_number(noise_floor_dbm, 1)} dBm`."
    )


def build_report_markdown(
    mobility: pd.DataFrame,
    fading_metrics: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    logical_channels: pd.DataFrame,
    rbw_noise: pd.DataFrame,
    red_checklist: pd.DataFrame,
) -> str:
    fc_mhz = float(mobility["carrier_frequency_mhz"].iloc[0])
    timeslot_ms = float(mobility["gsm_timeslot_ms"].iloc[0])
    timeslot_us = timeslot_ms * 1e3
    scenario_a_radius = float(mobility["cell_radius_km"].iloc[0])
    scenario_b_radius = float(frequency_planning["cell_radius_km"].iloc[0])
    cluster_size = int(frequency_planning["cluster_size_N"].iloc[0])
    total_carriers = int(frequency_planning["total_carriers"].iloc[0])
    carriers_per_cell = int(frequency_planning["carriers_per_cell"].iloc[0])
    reuse_ratio = float(frequency_planning["reuse_ratio_D_over_R"].iloc[0])
    reuse_distance_km = float(frequency_planning["reuse_distance_km"].iloc[0])
    speeds = ", ".join(f"{speed:g} km/h" for speed in mobility["speed_kmh"].tolist())
    bcch_cells = int(logical_channels[logical_channels["carrier_role"].str.contains("BCCH")]["cell"].nunique())
    best_stability = str(mobility.sort_values("coherence_to_timeslot_ratio", ascending=False)["stability_class"].iloc[0])
    worst_stability = str(mobility.sort_values("coherence_to_timeslot_ratio", ascending=True)["stability_class"].iloc[0])
    coherence_bw_khz = (
        float(mobility["coherence_bandwidth_khz"].iloc[0]) if "coherence_bandwidth_khz" in mobility.columns else None
    )
    estimated_cir_db = (
        float(frequency_planning["estimated_cir_db"].iloc[0]) if "estimated_cir_db" in frequency_planning.columns else None
    )
    displayed_noise_dbm = (
        float(rbw_noise["displayed_average_noise_dbm"].min()) if "displayed_average_noise_dbm" in rbw_noise.columns else None
    )
    coherence_bw_text = _format_number(coherence_bw_khz, 1) if coherence_bw_khz is not None else "n/d"
    estimated_cir_text = _format_number(estimated_cir_db, 2) if estimated_cir_db is not None else "n/d"
    displayed_noise_text = _format_number(displayed_noise_dbm, 1) if displayed_noise_dbm is not None else "n/d"

    lines: list[str] = []
    lines.append("# Protocolo Titán — Informe técnico docente")
    lines.append("")
    lines.append("## Resumen")
    lines.append("")
    lines.append(
        "Este informe evalúa la viabilidad de una red GSM/EDGE táctica en dos contextos de misión crítica: "
        "movilidad extrema en un convoy ferroviario y alta densidad operativa en un campamento base. "
        f"Se parte de una red GSM-900 con portadoras de 200 kHz, un timeslot de {_format_number(timeslot_us, 0)} µs, "
        f"velocidades de estudio {speeds} y un plan celular con {total_carriers} ARFCNs distribuidos en un clúster N={cluster_size}."
    )
    lines.append(
        "Los resultados muestran que el escenario A está dominado por la variación temporal del canal, "
        "mientras que el escenario B queda condicionado por la reutilización espectral, la protección co-canal y el criterio instrumental de medida."
    )
    if coherence_bw_khz is not None or estimated_cir_db is not None or displayed_noise_dbm is not None:
        lines.append(
            "Además del cálculo base, esta versión incorpora métricas de mayor realismo: "
            f"ancho de banda de coherencia del orden de {coherence_bw_text} kHz, "
            f"C/I estimada de {estimated_cir_text} dB en el campamento base y "
            f"DANL mostrado de hasta {displayed_noise_text} dBm en la configuración más sensible del analizador."
        )
    lines.append("")

    _append_section(
        lines,
        "## 1. Introducción",
        [
            (
                "El reto Protocolo Titán plantea el diseño de una red táctica de alta fiabilidad basada en GSM/EDGE. "
                "Aunque se trata de una arquitectura histórica, sigue siendo útil como referencia docente porque combina robustez, "
                "interoperabilidad, sencillez de despliegue y una estructura radio bien documentada [1], [2]."
            ),
            (
                "Desde la perspectiva de misión crítica, la red debe soportar coordinación de voz, control de equipos, telemetría de baja tasa "
                "y decisiones de ingeniería que preserven tanto la disponibilidad del enlace como el uso disciplinado del espectro."
            ),
            "### 1.1 Objetivo general",
            "",
            "Demostrar, con cálculos reproducibles y criterio ingenieril, la viabilidad de una red GSM/EDGE táctica en un escenario de alta movilidad y en un entorno de alta densidad operativa.",
            "",
            "### 1.2 Objetivos específicos",
            "",
            "- Cuantificar el efecto Doppler y el tiempo de coherencia para velocidades tácticas representativas.",
            "- Determinar si el canal puede considerarse aproximadamente estable durante un timeslot GSM.",
            "- Diseñar una planificación básica de frecuencias con 24 ARFCNs y clúster N=4.",
            "- Justificar el papel de BCCH, TCH y la política de potencia/hopping en el reparto de canales.",
            "- Analizar cómo cambia el suelo de ruido instrumental al modificar la RBW del analizador.",
            "- Integrar movilidad, reutilización y certificación RED en una propuesta coherente de despliegue.",
        ],
    )

    _append_section(
        lines,
        "## 2. Estado del arte",
        [
            "### 2.1 Acceso múltiple en GSM/EDGE",
            "",
            (
                "GSM organiza el acceso radio mediante FDMA y TDMA. En frecuencia, el espectro se divide en portadoras de 200 kHz; "
                f"en tiempo, cada portadora se reparte en 8 timeslots por trama, con un timeslot base de {_format_number(timeslot_us, 0)} µs [1], [2]."
            ),
            "",
            "### 2.2 Canales físicos y canales lógicos",
            "",
            (
                "Un canal físico es la combinación portadora-timeslot. Sobre él se multiplexan canales lógicos con finalidades distintas: "
                "BCCH para difusión y camping, TCH para tráfico y otros canales de control para señalización, acceso y sincronización. "
                "En consecuencia, la capacidad útil de usuario nunca debe interpretarse como una mera cuenta bruta de timeslots."
            ),
            "",
            "### 2.3 Fading, multitrayecto y movilidad",
            "",
            (
                "El fading por multitrayecto aparece cuando varias réplicas temporales de la señal interfieren constructiva o destructivamente. "
                "Un entorno sin línea de vista dominante suele aproximarse por Rayleigh, mientras que la presencia de una componente fuerte LOS "
                "conduce a un comportamiento Rician [3], [4]."
            ),
            (
                "Cuando el ancho de banda de la señal es pequeño frente al ancho de banda de coherencia, el canal puede modelarse como fading plano; "
                "si no, aparecen efectos selectivos en frecuencia. En este reto se usa una simulación docente de fading plano para estudiar la estabilidad intra-ráfaga."
            ),
            "",
            "### 2.4 Certificación radioeléctrica y Directiva RED",
            "",
            (
                "El despliegue radio no se limita a la capacidad de servicio: también exige control de emisiones no deseadas, "
                "uso eficiente del espectro y trazabilidad de medidas. La Directiva RED aporta el marco regulatorio de conformidad, "
                "mientras que el analizador de espectro y el ajuste de RBW permiten distinguir señales débiles de ruido instrumental [5]."
            ),
        ],
    )

    _append_section(
        lines,
        "## 3. Metodología",
        [
            "### 3.1 Hipótesis de partida",
            "",
            f"- Banda de operación: GSM-900 con frecuencia central aproximada de {fc_mhz:g} MHz.",
            "- Acceso radio: portadoras de 200 kHz y 8 timeslots por trama.",
            f"- Timeslot analizado: {_format_number(timeslot_us, 0)} µs.",
            "- Velocidad de propagación: `3 × 10^8 m/s`.",
            f"- Figura de ruido del analizador: {rbw_noise['noise_figure_db'].iloc[0]:g} dB.",
            "",
            "### 3.2 Ecuaciones empleadas",
            "",
            "- Conversión de velocidad: `v(m/s) = v(km/h) / 3.6`.",
            "- Doppler máximo: `f_d = v · f_c / c`.",
            "- Tiempo de coherencia: `T_c ≈ 0.423 / f_d`.",
            "- Ancho de banda de coherencia: `B_c ≈ 1 / (5τ_rms)`.",
            "- Reutilización celular: `D/R = √(3N)` y `D = R · √(3N)`.",
            "- Interferencia co-canal estimada: `C/I ≈ (D/R)^n / i0`.",
            "- Suelo de ruido instrumental: `N(dBm) = -174 + 10 log10(RBW) + NF`.",
            "- DANL mostrado: `N_DANL ≈ -174 + 10 log10(RBW·ENBW) + NF - bias_detector`.",
            "",
            "### 3.3 Flujo reproducible",
            "",
            "1. Convertir velocidades de km/h a m/s y calcular Doppler máximo.",
            "2. Obtener el tiempo de coherencia, el Doppler normalizado y el ancho de banda de coherencia.",
            "3. Simular fading Rayleigh y Rician con un proceso temporal correlacionado tipo Jakes multitrayecto.",
            "4. Repartir 24 portadoras entre las celdas del clúster y calcular la distancia de reutilización y la C/I.",
            "5. Evaluar el impacto de la RBW del analizador sobre el ruido integrado, el DANL y el tiempo de barrido.",
            "6. Integrar resultados en una discusión común orientada a decisiones de despliegue.",
        ],
    )

    _append_section(
        lines,
        "## 4. Escenarios de estudio",
        [
            "### 4.1 Escenario A — Convoy de alta velocidad",
            "",
            (
                f"Se estudia un despliegue lineal en un valle operativo con radio de celda {scenario_a_radius:g} km y velocidades {speeds}. "
                "Aquí domina la movilidad extrema: el problema principal no es la cantidad de portadoras disponibles, sino la rapidez con la que cambia el canal radio dentro de una ráfaga GSM."
            ),
            "",
            "### 4.2 Escenario B — Campamento base y certificación",
            "",
            (
                f"Se analiza un campamento base con radio de celda {scenario_b_radius:g} km, {total_carriers} portadoras disponibles y clúster N={cluster_size}. "
                "En este contexto dominan la planificación espectral, la interferencia co-canal y la verificación instrumental de emisiones."
            ),
            "",
            "Mini conclusión parcial: los dos escenarios son complementarios. El primero exige analizar estabilidad temporal del canal; el segundo exige disciplina de espectro y criterio de certificación.",
        ],
    )

    proof_lines = ["### 5.1 Prueba 1 — Movilidad y fading", ""]
    for _, row in mobility.iterrows():
        proof_lines.extend(_mobility_case_lines(row))
        proof_lines.append("")
    proof_lines.extend(
        [
            "Las métricas de fading exportadas permiten comparar la dispersión relativa de la envolvente en escenarios Rayleigh y Rician. "
            "En general, el caso Rician presenta menor variación pico a pico por la presencia de una componente dominante, mientras que la LCR y la AFD ofrecen una lectura más realista de la severidad de los fades.",
            "",
            "### 5.2 Prueba 2 — Planificación de frecuencias",
            "",
            f"1. Reparto espectral: `{total_carriers} / {cluster_size} = {carriers_per_cell}` portadoras por celda.",
            f"2. Relación de reutilización: `D/R = √(3N) = √(3 · {cluster_size}) = {_format_number(reuse_ratio)}`.",
            f"3. Distancia de protección: `D = R · √(3N) = {scenario_b_radius:g} · {_format_number(reuse_ratio)} = {_format_number(reuse_distance_km)} km`.",
            (
                f"4. C/I estimada en borde de celda: `10 log10((D/R)^n / i0) = {_format_number(estimated_cir_db, 2)} dB`."
                if estimated_cir_db is not None
                else "4. C/I estimada en borde de celda: se obtiene a partir del patrón de reutilización y del exponente de pérdidas."
            ),
            (
                "5. Interpretación: la distancia de reutilización actúa como separación geométrica mínima entre celdas co-canal; "
                "cuanto mayor es, mayor protección frente a interferencia, pero menor densidad espectral efectiva."
            ),
            (
                "6. Política BCCH/TCH: BCCH se mantiene estable y a potencia fija para asegurar camping, sincronización y acceso; "
                "en cambio, TCH puede beneficiarse de hopping y control de potencia cuando la infraestructura lo soporte."
            ),
            "",
            "Mini conclusión parcial: el clúster N=4 introduce un compromiso razonable entre capacidad y protección co-canal para un despliegue táctico con disciplina espectral.",
            "",
            "### 5.3 Prueba 3 — Certificación y suelo de ruido",
            "",
        ]
    )
    proof_lines.extend(_noise_case_lines(row) for _, row in rbw_noise.iterrows())
    proof_lines.extend(
        [
            "",
            (
                "Interpretación: al reducir la RBW disminuye el ruido integrado por el analizador, lo que mejora la visibilidad de señales débiles. "
                "Sin embargo, esa mejora no implica que la señal real aumente: simplemente mejora la relación visual señal-ruido en la medida."
            ),
            (
                "Compromiso experimental: una RBW muy estrecha reduce el ruido mostrado, pero incrementa el tiempo de barrido y puede ralentizar la validación de conformidad. "
                "La corrección por ENBW y sesgo del detector acerca mejor la tabla a lo que realmente enseña el analizador en pantalla."
            ),
            "",
            "Mini conclusión parcial: el analizador debe ajustarse con criterio. Una RBW demasiado ancha oculta señales débiles en el ruido; una demasiado estrecha penaliza el tiempo de medida.",
            "",
            "### 5.4 Interpretación integrada",
            "",
            (
                "Los resultados no deben leerse como tres bloques independientes. El escenario A demuestra que la estabilidad intra-ráfaga puede mantenerse incluso con movilidad elevada, "
                "lo que avala el uso de GSM/EDGE cuando se prioriza robustez y simplicidad. El escenario B muestra que esa robustez solo es útil si la planificación de frecuencias, "
                "el reparto BCCH/TCH y el criterio instrumental preservan el servicio frente a interferencia y falsas interpretaciones de laboratorio."
            ),
        ]
    )
    _append_section(lines, "## 5. Pruebas, cálculos y simulaciones", proof_lines)

    _append_section(
        lines,
        "## 6. Resultados",
        [
            "### 6.1 Tabla comparativa de movilidad",
            "",
            _dataframe_to_markdown(mobility),
            "",
            "### 6.2 Métricas de fading",
            "",
            _dataframe_to_markdown(fading_metrics),
            "",
            "### 6.3 Planificación de frecuencias",
            "",
            _dataframe_to_markdown(frequency_planning),
            "",
            "### 6.4 Canales lógicos y físicos",
            "",
            _dataframe_to_markdown(logical_channels),
            "",
            "### 6.5 Instrumentación y RBW",
            "",
            _dataframe_to_markdown(rbw_noise),
            "",
            "### 6.6 Checklist RED orientativo",
            "",
            _dataframe_to_markdown(red_checklist),
            "",
            "### 6.7 Figuras recomendadas para el informe",
            "",
            "![Figura 1. Doppler máximo frente a velocidad](figures/escenario_a_doppler.png)",
            "",
            "![Figura 2. Coherencia frente al timeslot GSM](figures/escenario_a_coherencia_vs_timeslot.png)",
            "",
            "![Figura 3. Reutilización celular en el campamento base](figures/escenario_b_reutilizacion.png)",
            "",
            "![Figura 4. Mapa de clúster y patrón de reutilización](figures/escenario_b_cluster_map.png)",
            "",
            "![Figura 5. Distribución de portadoras por celda](figures/escenario_b_distribucion_portadoras.png)",
            "",
            "![Figura 6. Huella espectral aproximada del despliegue](figures/escenario_b_spectrum.png)",
            "",
            "![Figura 7. Ruido instrumental frente a RBW](figures/certificacion_rbw_ruido.png)",
        ],
    )

    _append_section(
        lines,
        "## 7. Discusión y conclusiones",
        [
            "### 7.1 Discusión",
            "",
            (
                f"En el escenario A, el peor caso temporal sigue siendo **{worst_stability}**, mientras que el mejor caso alcanza **{best_stability}**. "
                "Esto sugiere que la red puede operar con márgenes razonables en movilidad, siempre que el diseño conserve codificación robusta, control de potencia y suficiente sensibilidad radio."
            ),
            (
                f"En el escenario B, el reparto de {carriers_per_cell} portadoras por celda y la distancia de protección de {_format_number(reuse_distance_km)} km "
                "ofrecen una configuración coherente con un despliegue disciplinado. La existencia de "
                f"{bcch_cells} portadoras BCCH dedicadas refuerza la accesibilidad de la red a costa de reducir ligeramente el tráfico bruto disponible."
            ),
            (
                "Desde el punto de vista de certificación, la interpretación correcta de la RBW evita errores frecuentes: reducir ruido mostrado no significa aumentar potencia real. "
                "Ese matiz es crucial para no sobreestimar el comportamiento de emisiones espurias o señales débiles."
            ),
            "",
            "### 7.2 Conclusiones",
            "",
            "1. GSM/EDGE sigue siendo una referencia válida para entornos tácticos donde priman robustez, interoperabilidad y sencillez operativa.",
            "2. La movilidad analizada no invalida el enlace intra-ráfaga; el impacto de Doppler es apreciable pero gestionable con diseño conservador.",
            f"3. El clúster N={cluster_size} proporciona un equilibrio razonable entre capacidad y protección co-canal en el campamento base.",
            "4. BCCH debe mantenerse estable y fácilmente detectable; TCH puede optimizarse con mayor flexibilidad operativa.",
            "5. La configuración del analizador de espectro forma parte del diseño ingenieril porque condiciona la lectura correcta de conformidad y emisiones no deseadas.",
        ],
    )

    _append_section(
        lines,
        "## 8. Limitaciones y trabajo futuro",
        [
            "- El modelo de fading empleado es docente y no sustituye a una simulación completa tipo Jakes o a medidas de campo.",
            "- La asignación de ARFCNs es consecutiva y no incorpora coordinación real basada en campañas de medida.",
            "- La capacidad de tráfico se resume de forma pedagógica y no modela todos los canales lógicos de control de una red comercial completa.",
            "- Como trabajo futuro, puede ampliarse el proyecto con balance de enlace, C/I estimada y comparación con tecnologías más modernas de misión crítica.",
        ],
    )

    _append_section(
        lines,
        "## 9. Referencias",
        [
            "[1] 3GPP TS 45.002, *Multiplexing and multiple access on the radio path*.",
            "[2] 3GPP TS 45.001, *Physical layer on the radio path; General description*.",
            "[3] T. S. Rappaport, *Wireless Communications: Principles and Practice*, 2nd ed. Prentice Hall, 2002.",
            "[4] B. Sklar, *Digital Communications: Fundamentals and Applications*, 2nd ed. Prentice Hall, 2001.",
            "[5] Directive 2014/53/EU of the European Parliament and of the Council of 16 April 2014 on radio equipment (RED).",
        ],
    )

    return "\n".join(lines).rstrip() + "\n"


def build_calculation_annex_markdown(
    mobility: pd.DataFrame,
    fading_metrics: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    logical_channels: pd.DataFrame,
    rbw_noise: pd.DataFrame,
    red_checklist: pd.DataFrame,
) -> str:
    lines = [
        "# Anexo de cálculos y trazabilidad",
        "",
        "## A.1 Parámetros base utilizados",
        "",
        _dataframe_to_markdown(mobility[["cell_radius_km", "carrier_frequency_mhz", "gsm_timeslot_ms"]].drop_duplicates()),
        "",
        "## A.2 Desarrollo detallado del escenario A",
        "",
    ]

    for _, row in mobility.iterrows():
        lines.extend(_mobility_case_lines(row))
        lines.append("")

    lines.extend(
        [
            "## A.3 Métricas de fading exportadas",
            "",
            _dataframe_to_markdown(fading_metrics),
            "",
            "## A.4 Planificación celular y asignación de ARFCNs",
            "",
            _dataframe_to_markdown(frequency_planning),
            "",
            _dataframe_to_markdown(logical_channels),
            "",
            "## A.5 Instrumentación y checklist RED",
            "",
            _dataframe_to_markdown(rbw_noise),
            "",
        ]
    )
    lines.extend(_noise_case_lines(row) for _, row in rbw_noise.iterrows())
    lines.extend(
        [
            "",
            _dataframe_to_markdown(red_checklist),
            "",
            "## A.6 Artefactos generados",
            "",
            "- `outputs/escenario_a_movilidad.csv` y `outputs/escenario_a_fading_metricas.csv`.",
            "- `outputs/escenario_b_planificacion.csv` y `outputs/escenario_b_canales_logicos.csv`.",
            "- `outputs/certificacion_rbw.csv` y `outputs/certificacion_red_checklist.csv`.",
            "- Trazas `outputs/traza_*.csv` y figuras PNG en `outputs/figures/`.",
        ]
    )
    return "\n".join(lines).rstrip() + "\n"


def build_defense_brief_markdown(
    mobility: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    rbw_noise: pd.DataFrame,
) -> str:
    worst_case = mobility.sort_values("coherence_to_timeslot_ratio").iloc[0]
    reuse_distance_km = float(frequency_planning["reuse_distance_km"].iloc[0])
    cluster_size = int(frequency_planning["cluster_size_N"].iloc[0])
    noise_best = float(rbw_noise.sort_values("noise_floor_dbm").iloc[0]["noise_floor_dbm"])
    return "\n".join(
        [
            "# Guion breve de defensa",
            "",
            "## Diapositiva 1 — Misión y objetivo",
            "- Problema: desplegar una red GSM/EDGE táctica para movilidad extrema y alta densidad operativa.",
            "- Objetivo: demostrar viabilidad técnica con cálculos reproducibles y criterio profesional.",
            "",
            "## Diapositiva 2 — Marco teórico esencial",
            "- FDMA de 200 kHz y TDMA de 8 timeslots por trama.",
            "- Diferencia entre canales físicos y lógicos, con foco en BCCH y TCH.",
            "- Fading Rayleigh/Rician y criterio RED de conformidad espectral.",
            "",
            "## Diapositiva 3 — Escenario A",
            f"- Peor caso estudiado: {worst_case['speed_kmh']:g} km/h con clase de estabilidad {worst_case['stability_class']}.",
            "- Mensaje clave: la movilidad reduce el tiempo de coherencia, pero no invalida el enlace intra-ráfaga con los parámetros base.",
            "",
            "## Diapositiva 4 — Escenario B",
            f"- Clúster N={cluster_size} y distancia de reutilización de {_format_number(reuse_distance_km)} km.",
            "- Mensaje clave: la disciplina espectral protege frente a co-canal sin sacrificar totalmente la capacidad táctica.",
            "",
            "## Diapositiva 5 — Instrumentación",
            f"- Mejor suelo de ruido mostrado: {_format_number(noise_best, 1)} dBm con la RBW más estrecha.",
            "- Mensaje clave: una RBW menor reduce ruido integrado, pero alarga la medida.",
            "",
            "## Diapositiva 6 — Conclusión final",
            "- El proyecto conecta teoría, cálculo y criterio de despliegue.",
            "- La propuesta es defendible para misión crítica si se mantiene diseño conservador, planificación ordenada y validación instrumental rigurosa.",
        ]
    ) + "\n"


def _markdown_to_html_fragment(markdown_text: str) -> str:
    return markdown.markdown(markdown_text, extensions=["tables", "sane_lists"])


def build_html_document(document_title: str, markdown_text: str) -> str:
    body = _markdown_to_html_fragment(markdown_text)
    title = html_escape(document_title)
    return f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    :root {{ color-scheme: light; }}
    body {{
      font-family: Arial, Helvetica, sans-serif;
      margin: 0;
      background: #f5f7fb;
      color: #1f2937;
      line-height: 1.55;
    }}
    main {{
      max-width: 1080px;
      margin: 32px auto;
      background: #ffffff;
      padding: 36px 42px;
      box-shadow: 0 8px 24px rgba(15, 23, 42, 0.08);
      border-radius: 14px;
    }}
    h1, h2, h3, h4 {{ color: #0f172a; margin-top: 1.35em; }}
    h1 {{ font-size: 2rem; border-bottom: 2px solid #cbd5e1; padding-bottom: 0.35rem; }}
    h2 {{ font-size: 1.5rem; border-bottom: 1px solid #e2e8f0; padding-bottom: 0.25rem; }}
    h3 {{ font-size: 1.2rem; }}
    h4 {{ font-size: 1.05rem; }}
    p, li {{ font-size: 0.98rem; }}
    code {{ background: #eef2ff; padding: 0.1rem 0.3rem; border-radius: 4px; font-family: Consolas, monospace; }}
    pre {{ background: #0f172a; color: #f8fafc; padding: 1rem; overflow-x: auto; border-radius: 8px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 1rem 0 1.4rem 0; font-size: 0.92rem; }}
    th, td {{ border: 1px solid #cbd5e1; padding: 0.55rem 0.6rem; text-align: left; vertical-align: top; }}
    th {{ background: #e2e8f0; }}
    img {{ max-width: 100%; height: auto; display: block; margin: 1rem auto; border: 1px solid #cbd5e1; border-radius: 8px; }}
    blockquote {{ margin: 1rem 0; padding-left: 1rem; border-left: 4px solid #94a3b8; color: #475569; }}
    @media print {{
      body {{ background: #ffffff; }}
      main {{ box-shadow: none; margin: 0; max-width: none; padding: 0; }}
      a {{ color: inherit; text-decoration: none; }}
    }}
  </style>
</head>
<body>
  <main>
{body}
  </main>
</body>
</html>
"""


def _inline_markdown_to_reportlab(text: str) -> str:
    escaped = html_escape(text, quote=True)
    escaped = re.sub(r"`([^`]+)`", lambda match: f'<font name="Courier">{match.group(1)}</font>', escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", escaped)
    return escaped


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"(`|\*\*)", "", text).replace("*", "")


def _add_inline_runs(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))")
    last_idx = 0

    for match in token_pattern.finditer(text):
        if match.start() > last_idx:
            paragraph.add_run(text[last_idx:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True

        last_idx = match.end()

    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])


def _add_docx_cover_page(document: Document, report_title: str, document_title: str) -> None:
    today = date.today().strftime("%d/%m/%Y")

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("PLANTILLA DE ENTREGA ACADÉMICA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)

    document.add_paragraph()
    document.add_paragraph()

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(report_title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(document_title)
    subtitle_run.italic = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)

    summary_paragraph = document.add_paragraph()
    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_paragraph.add_run(
        "Análisis técnico de una red GSM/EDGE táctica para movilidad extrema, planificación espectral y certificación instrumental."
    )
    summary_run.font.name = "Arial"
    summary_run.font.size = Pt(10)

    document.add_paragraph()
    document.add_paragraph()

    meta_table = document.add_table(rows=7, cols=2)
    meta_table.style = "Table Grid"
    metadata_rows = [
        ("Centro / Universidad", "[Completar]"),
        ("Asignatura / Módulo", "[Completar]"),
        ("Titulación / Grupo", "[Completar]"),
        ("Autor / Equipo", "[Completar]"),
        ("Docente", "[Completar]"),
        ("Fecha de entrega", today),
        ("Versión del documento", "v1.0"),
    ]
    for row_idx, (label, value) in enumerate(metadata_rows):
        label_cell = meta_table.cell(row_idx, 0).paragraphs[0]
        value_cell = meta_table.cell(row_idx, 1).paragraphs[0]
        label_run = label_cell.add_run(label)
        label_run.bold = True
        label_run.font.name = "Arial"
        value_run = value_cell.add_run(value)
        value_run.font.name = "Arial"

    document.add_paragraph()

    body_title = document.add_paragraph()
    body_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body_title_run = body_title.add_run("Datos de entrega y revisión final")
    body_title_run.bold = True
    body_title_run.font.name = "Arial"
    body_title_run.font.size = Pt(11)

    checklist_items = [
        "Revisar portada, índice, numeración y formato de referencias antes de exportar la versión final.",
        "Verificar que las figuras insertadas mantengan legibilidad en PDF y en impresión.",
        "Sustituir los campos [Completar] por los datos reales del grupo o del estudiante.",
    ]
    for item in checklist_items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Arial"

    document.add_page_break()


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    return all(set(cell.strip()) <= {"-", ":"} and cell.strip() for cell in stripped.split("|"))


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _markdown_to_pdf_story(markdown_text: str, asset_root: Path) -> list:
    stylesheet = getSampleStyleSheet()
    styles = {
        "title": stylesheet["Title"],
        "h1": stylesheet["Heading1"],
        "h2": stylesheet["Heading2"],
        "h3": stylesheet["Heading3"],
        "h4": ParagraphStyle("Heading4Custom", parent=stylesheet["Heading3"], fontSize=11, leading=14, spaceAfter=6),
        "body": ParagraphStyle("BodyCustom", parent=stylesheet["BodyText"], fontSize=9.5, leading=13, spaceAfter=6),
        "caption": ParagraphStyle("CaptionCustom", parent=stylesheet["BodyText"], fontSize=8.5, leading=11, alignment=1, textColor=colors.HexColor("#334155")),
    }

    lines = markdown_text.splitlines()
    story: list = []
    idx = 0

    while idx < len(lines):
        line = lines[idx].rstrip()

        if not line.strip():
            idx += 1
            continue

        if line.startswith("# "):
            story.append(Paragraph(_inline_markdown_to_reportlab(line[2:].strip()), styles["title"]))
            story.append(Spacer(1, 0.2 * cm))
            idx += 1
            continue

        if line.startswith("## "):
            story.append(Paragraph(_inline_markdown_to_reportlab(line[3:].strip()), styles["h1"]))
            idx += 1
            continue

        if line.startswith("### "):
            story.append(Paragraph(_inline_markdown_to_reportlab(line[4:].strip()), styles["h2"]))
            idx += 1
            continue

        if line.startswith("#### "):
            story.append(Paragraph(_inline_markdown_to_reportlab(line[5:].strip()), styles["h4"]))
            idx += 1
            continue

        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, relative_path = match.groups()
                image_path = asset_root / relative_path
                if image_path.exists():
                    image = Image(str(image_path))
                    image._restrictSize(25 * cm, 14 * cm)
                    story.append(image)
                    if caption:
                        story.append(Paragraph(_inline_markdown_to_reportlab(caption), styles["caption"]))
                    story.append(Spacer(1, 0.2 * cm))
            idx += 1
            continue

        if line.startswith("|"):
            table_lines = [line]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].rstrip())
                idx += 1

            parsed_rows = [_parse_table_row(table_line) for table_line in table_lines if not _is_table_separator(table_line)]
            if parsed_rows:
                col_count = max(len(row) for row in parsed_rows)
                normalized_rows = [
                    row + [""] * (col_count - len(row)) for row in parsed_rows
                ]
                table_data = [
                    [Paragraph(_inline_markdown_to_reportlab(cell), styles["body"]) for cell in row]
                    for row in normalized_rows
                ]
                col_width = 27.0 * cm / col_count
                table = Table(table_data, repeatRows=1, colWidths=[col_width] * col_count)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 0.25 * cm))
            continue

        if line.startswith("- "):
            list_lines = []
            while idx < len(lines) and lines[idx].startswith("- "):
                list_lines.append(lines[idx][2:].strip())
                idx += 1
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_markdown_to_reportlab(item), styles["body"])) for item in list_lines],
                    bulletType="bullet",
                    leftIndent=14,
                )
            )
            story.append(Spacer(1, 0.15 * cm))
            continue

        if re.match(r"^\d+\.\s", line):
            list_lines = []
            while idx < len(lines) and re.match(r"^\d+\.\s", lines[idx]):
                list_lines.append(re.sub(r"^\d+\.\s", "", lines[idx]).strip())
                idx += 1
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_markdown_to_reportlab(item), styles["body"])) for item in list_lines],
                    bulletType="1",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 0.15 * cm))
            continue

        paragraph_lines = [line]
        idx += 1
        while idx < len(lines):
            next_line = lines[idx].rstrip()
            if not next_line.strip():
                break
            if next_line.startswith(("#", "!", "|", "- ")) or re.match(r"^\d+\.\s", next_line):
                break
            paragraph_lines.append(next_line)
            idx += 1

        paragraph_text = " ".join(part.strip() for part in paragraph_lines)
        story.append(Paragraph(_inline_markdown_to_reportlab(paragraph_text), styles["body"]))

    return story


def _draw_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 0.6 * cm, f"Página {canvas.getPageNumber()}")
    canvas.restoreState()


def write_html_document(output_path: Path, document_title: str, markdown_text: str) -> None:
    output_path.write_text(build_html_document(document_title, markdown_text), encoding="utf-8")


def write_pdf_document(output_path: Path, document_title: str, markdown_text: str, asset_root: Path) -> None:
    doc = SimpleDocTemplate(
        str(output_path),
        title=document_title,
        pagesize=landscape(A4),
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    story = _markdown_to_pdf_story(markdown_text, asset_root=asset_root)
    doc.build(story, onFirstPage=_draw_page_number, onLaterPages=_draw_page_number)


def write_docx_document(output_path: Path, document_title: str, markdown_text: str, asset_root: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    heading_styles = {
        1: document.styles["Heading 1"],
        2: document.styles["Heading 2"],
        3: document.styles["Heading 3"],
    }
    for style in heading_styles.values():
        style.font.name = "Arial"

    lines = markdown_text.splitlines()
    idx = 0
    cover_added = False

    while idx < len(lines):
        line = lines[idx].rstrip()

        if not line.strip():
            idx += 1
            continue

        if line.startswith("# "):
            if not cover_added:
                _add_docx_cover_page(document, line[2:].strip(), document_title)
                cover_added = True
            idx += 1
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            _add_inline_runs(paragraph, line[3:].strip())
            idx += 1
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            _add_inline_runs(paragraph, line[4:].strip())
            idx += 1
            continue

        if line.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            _add_inline_runs(paragraph, line[5:].strip())
            idx += 1
            continue

        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, relative_path = match.groups()
                image_path = asset_root / relative_path
                if image_path.exists():
                    document.add_picture(str(image_path), width=Cm(16.5))
                    caption_paragraph = document.add_paragraph()
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_paragraph.add_run(caption)
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
            idx += 1
            continue

        if line.startswith("|"):
            table_lines = [line]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].rstrip())
                idx += 1

            parsed_rows = [_parse_table_row(table_line) for table_line in table_lines if not _is_table_separator(table_line)]
            if parsed_rows:
                col_count = max(len(row) for row in parsed_rows)
                table = document.add_table(rows=len(parsed_rows), cols=col_count)
                table.style = "Table Grid"
                for row_idx, row in enumerate(parsed_rows):
                    for col_idx in range(col_count):
                        value = row[col_idx] if col_idx < len(row) else ""
                        cell_paragraph = table.cell(row_idx, col_idx).paragraphs[0]
                        _add_inline_runs(cell_paragraph, _strip_inline_markdown(value))
                        if row_idx == 0:
                            for run in cell_paragraph.runs:
                                run.bold = True
            continue

        if line.startswith("- "):
            while idx < len(lines) and lines[idx].startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                _add_inline_runs(paragraph, lines[idx][2:].strip())
                idx += 1
            continue

        if re.match(r"^\d+\.\s", line):
            while idx < len(lines) and re.match(r"^\d+\.\s", lines[idx]):
                paragraph = document.add_paragraph(style="List Number")
                _add_inline_runs(paragraph, re.sub(r"^\d+\.\s", "", lines[idx]).strip())
                idx += 1
            continue

        paragraph_lines = [line]
        idx += 1
        while idx < len(lines):
            next_line = lines[idx].rstrip()
            if not next_line.strip():
                break
            if next_line.startswith(("#", "!", "|", "- ")) or re.match(r"^\d+\.\s", next_line):
                break
            paragraph_lines.append(next_line)
            idx += 1

        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, " ".join(part.strip() for part in paragraph_lines))

    document.save(str(output_path))


def write_markdown_report(
    output_path: Path,
    mobility: pd.DataFrame,
    fading_metrics: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    logical_channels: pd.DataFrame,
    rbw_noise: pd.DataFrame,
    red_checklist: pd.DataFrame,
) -> None:
    """Genera un informe docente-técnico alineado con la rúbrica del reto."""
    output_path.write_text(
        build_report_markdown(
            mobility=mobility,
            fading_metrics=fading_metrics,
            frequency_planning=frequency_planning,
            logical_channels=logical_channels,
            rbw_noise=rbw_noise,
            red_checklist=red_checklist,
        ),
        encoding="utf-8",
    )


def write_calculation_annex(
    output_path: Path,
    mobility: pd.DataFrame,
    fading_metrics: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    logical_channels: pd.DataFrame,
    rbw_noise: pd.DataFrame,
    red_checklist: pd.DataFrame,
) -> None:
    output_path.write_text(
        build_calculation_annex_markdown(
            mobility=mobility,
            fading_metrics=fading_metrics,
            frequency_planning=frequency_planning,
            logical_channels=logical_channels,
            rbw_noise=rbw_noise,
            red_checklist=red_checklist,
        ),
        encoding="utf-8",
    )


def write_defense_brief(
    output_path: Path,
    mobility: pd.DataFrame,
    frequency_planning: pd.DataFrame,
    rbw_noise: pd.DataFrame,
) -> None:
    output_path.write_text(
        build_defense_brief_markdown(
            mobility=mobility,
            frequency_planning=frequency_planning,
            rbw_noise=rbw_noise,
        ),
        encoding="utf-8",
    )
